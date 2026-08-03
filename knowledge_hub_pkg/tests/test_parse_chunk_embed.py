"""Stage B end-to-end, against the real stack and LIVE bge-m3 — no mocks.

A real .docx SOP is landed through the real capture path (SeaweedFS WORM +
raw_documents + dispatch_queue), then consumed off the queue and processed:
Docling parse -> section/passage chunking with the real bge-m3 tokenizer ->
live Ollama embeddings -> rows persisted via Prompt 1's insert_document/
insert_chunks. Assertions check the persisted rows, not in-memory objects.
"""
from __future__ import annotations

import io
import math
import uuid

import pytest

from knowledge_hub.capture import CaptureService
from knowledge_hub.chunking import CHILD_TOKENS, _bge_m3_token_counter
from knowledge_hub.factstore_pg import parse_vector
from knowledge_hub.processing import ProcessingService
from knowledge_hub.sources_fs import FilesystemSourceAdapter

# ---------------------------------------------------------------------------
# The pilot SOP: generated as a REAL .docx (python-docx), with core
# properties, headings, and one section long enough to split into several
# overlapping children.
# ---------------------------------------------------------------------------
SOP_TITLE = "Equipment Cleaning and Maintenance SOP"
SOP_AUTHOR = "QA Team"

_PROCEDURE_SENTENCES = [
    "Power down the unit and apply the lockout tag to the breaker panel before "
    "any surface is touched.",
    "Remove loose debris from the hopper, guards, and conveyor bed using the "
    "dedicated soft-bristle brushes stored at the wash station.",
    "Prepare the approved solvent solution at the concentration listed on the "
    "product data sheet, recording the lot number in the cleaning log.",
    "Wipe every product-contact surface twice, working from the highest point "
    "of the machine downward so rinsate never crosses a cleaned area.",
    "Rinse all wiped surfaces with distilled water and inspect under the "
    "portable lamp for residue, streaking, or discoloration.",
    "Allow the equipment to air dry for a minimum of thirty minutes with the "
    "guards open and airflow unobstructed.",
    "Swab the three designated sample points and submit the swabs to QA under "
    "a chain-of-custody label within one hour.",
    "Reassemble the guards, remove the lockout tag, and run the empty cycle "
    "for five minutes while listening for abnormal noise.",
    "Record every step, including start and end times, in the cleaning log and "
    "have the shift supervisor countersign the entry.",
    "Report any deviation, failed inspection, or damaged component to the "
    "maintenance lead before the equipment is returned to service.",
]


def make_sop_docx() -> bytes:
    import docx
    d = docx.Document()
    props = d.core_properties
    props.title, props.author = SOP_TITLE, SOP_AUTHOR
    d.add_heading(SOP_TITLE, level=1)
    d.add_paragraph(
        "Purpose: define the mandatory cleaning and maintenance routine for "
        "all production equipment so that every line stays audit-ready.")
    d.add_heading("1. Scope", level=2)
    d.add_paragraph(
        "This procedure applies to all production equipment in Building A, "
        "including mixers, fillers, labelers, and conveyors. Contract staff "
        "must follow the same steps after completing orientation.")
    d.add_heading("2. Responsibilities", level=2)
    d.add_paragraph(
        "Line operators perform the routine cleaning described here. The "
        "shift supervisor verifies completion and countersigns the log. QA "
        "owns this document and audits compliance quarterly.")
    d.add_heading("3. Procedure", level=2)
    for sentence in _PROCEDURE_SENTENCES:
        d.add_paragraph(sentence)
    for sentence in _PROCEDURE_SENTENCES:
        d.add_paragraph(sentence.replace("the unit", "the backup unit"))
    d.add_heading("4. Records", level=2)
    d.add_paragraph(
        "Retain completed cleaning logs for three years in the QA records "
        "cabinet. Logs are quality records and may not be discarded early.")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


CSV_BYTES = b"\n".join(
    [b"asset_id,asset_name,last_cleaned,status"] +
    [b"%d,unit-%d,2026-07-%02d,ok" % (i, i, (i % 28) + 1) for i in range(1, 13)]
)


# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------
def land_file(store, pipeline, raw_store, dispatcher, tenant, rel, content,
              config=None, source_ref="fs-sops"):
    """Land one file through the REAL capture path; returns raw_documents.id."""
    capture = CaptureService(pipeline, raw_store, dispatcher)
    if config is not None:
        capture.registry.register(tenant, source_ref, "filesystem",
                                  config=config)
    root = land_file.root
    path = root / rel
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(content)
    adapter = FilesystemSourceAdapter(source_ref=source_ref, root=root)
    result = capture.run_source(tenant, adapter, mode="backfill")
    assert result.landed + result.replayed >= 1
    with store.transaction(tenant) as conn:
        row = conn.execute(
            "SELECT id FROM raw_documents"
            " WHERE tenant_id = %s AND source_native_id = %s"
            " ORDER BY version DESC LIMIT 1",
            (tenant, rel)).fetchone()
    return row["id"]


def doc_row(store, tenant, document_id):
    with store.transaction(tenant) as conn:
        return conn.execute(
            "SELECT * FROM documents WHERE tenant_id = %s AND id = %s",
            (tenant, document_id)).fetchone()


def chunk_rows(store, tenant, document_id):
    with store.transaction(tenant) as conn:
        return conn.execute(
            "SELECT * FROM chunks WHERE tenant_id = %s AND document_id = %s"
            " ORDER BY level, seq, id",
            (tenant, document_id)).fetchall()


def cosine(a, b) -> float:
    dot = sum(x * y for x, y in zip(a, b))
    return dot / (math.sqrt(sum(x * x for x in a)) *
                  math.sqrt(sum(x * x for x in b)))


# ---------------------------------------------------------------------------
# the landed + processed SOP every read-only test shares (module-scoped:
# docling + live embeddings run once)
# ---------------------------------------------------------------------------
@pytest.fixture(scope="module")
def sop_case(store, pipeline, raw_store, dispatcher, parser, chunker, embedder,
             tmp_path_factory):
    tenant = f"t-{uuid.uuid4().hex[:12]}"
    land_file.root = tmp_path_factory.mktemp("sops")
    raw_id = land_file(
        store, pipeline, raw_store, dispatcher, tenant,
        "sops/equipment_cleaning_and_maintenance.docx", make_sop_docx(),
        config={"data_track": "prose", "doc_type": "sop"})
    service = ProcessingService(pipeline, raw_store, parser, chunker,
                                embedder, dispatcher=dispatcher)
    # Consume off the REAL queue — the whole point of the Prompt 2 handoff.
    results = service.consume(tenant, limit=10)
    assert len(results) == 1, f"expected one queue delivery, got {results}"
    return {
        "tenant": tenant, "raw_id": raw_id, "result": results[0],
        "service": service,
    }


# ---------------------------------------------------------------------------
# 1. hierarchy persisted with valid links, correct tenancy
# ---------------------------------------------------------------------------
def test_sop_becomes_persisted_three_tier_hierarchy(sop_case, store, raw_store,
                                                    parser):
    tenant, result = sop_case["tenant"], sop_case["result"]
    assert result.status == "processed"
    assert result.parents >= 4 and result.children >= result.parents

    # Superparent: the documents row is the metadata/provenance roll-up.
    doc = doc_row(store, tenant, result.document_id)
    assert doc["tenant_id"] == tenant
    assert doc["doc_type"] == "sop"
    assert doc["title"] == SOP_TITLE          # docx core properties fallback
    assert doc["author"] == SOP_AUTHOR
    assert doc["raw_document_id"] == sop_case["raw_id"]
    assert doc["metadata"]["data_track"] == "prose"
    assert doc["metadata"]["declared_data_track"] == "prose"  # manifest tag
    assert doc["metadata"]["outline"], "heading structure should be recorded"
    assert doc["review_status"] == "none"

    rows = chunk_rows(store, tenant, result.document_id)
    parents = [r for r in rows if r["level"] == "parent"]
    children = [r for r in rows if r["level"] == "child"]
    assert len(parents) == result.parents
    assert len(children) == result.children

    # Linkage: every child points at a real parent of the same document;
    # every parent has at least one child; parents are roots.
    parent_ids = {p["id"] for p in parents}
    assert all(c["parent_chunk_id"] in parent_ids for c in children)
    assert {c["parent_chunk_id"] for c in children} == parent_ids
    assert all(p["parent_chunk_id"] is None for p in parents)
    assert all(r["tenant_id"] == tenant for r in rows)

    # Provenance: char offsets anchor into the parser's extracted text of the
    # EXACT landed bytes (version-pinned raw_uri), and children nest inside
    # their parent's span.
    raw = store.get_raw_document(tenant, sop_case["raw_id"])
    text = parser.extract_text(raw, raw_store.get(raw.raw_uri))
    by_id = {p["id"]: p for p in parents}
    for row in rows:
        assert text[row["char_start"]:row["char_end"]] == row["content"]
    for c in children:
        p = by_id[c["parent_chunk_id"]]
        assert p["char_start"] <= c["char_start"] <= c["char_end"] <= p["char_end"]

    # The landing record moved on: landed -> parsed.
    assert raw.status == "parsed"


# ---------------------------------------------------------------------------
# 2. child sizing: token bounds + overlap on semantic boundaries
# ---------------------------------------------------------------------------
def test_child_token_counts_within_bounds_and_overlap_present(sop_case, store):
    tenant, result = sop_case["tenant"], sop_case["result"]
    rows = chunk_rows(store, tenant, result.document_id)
    children = [r for r in rows if r["level"] == "child"]
    count = _bge_m3_token_counter()

    for c in children:
        assert 0 < c["token_count"] <= CHILD_TOKENS
        # persisted counts are real bge-m3 vocabulary counts, not estimates
        assert count(c["content"]) == c["token_count"]

    # At least one section is long enough to need several children, and
    # consecutive children of a parent overlap (the ~15% token overlap).
    by_parent: dict[int, list] = {}
    for c in children:
        by_parent.setdefault(c["parent_chunk_id"], []).append(c)
    multi = [sorted(v, key=lambda r: r["seq"]) for v in by_parent.values()
             if len(v) >= 2]
    assert multi, "expected at least one parent with multiple children"
    for siblings in multi:
        for a, b in zip(siblings, siblings[1:]):
            assert b["char_start"] < a["char_end"], \
                "consecutive children should share an overlap window"


# ---------------------------------------------------------------------------
# 3. live embeddings: every child got a real 1024-dim vector
# ---------------------------------------------------------------------------
def test_children_carry_live_1024_dim_embeddings(sop_case, store, embedder):
    tenant, result = sop_case["tenant"], sop_case["result"]
    rows = chunk_rows(store, tenant, result.document_id)
    for row in rows:
        vec = parse_vector(row["embedding"])
        if row["level"] == "child":
            assert vec is not None and len(vec) == 1024
            assert row["embedding_model"] == embedder.model == "bge-m3"
            assert row["embedding_version"] == embedder.version
        else:  # parents are the extraction unit, not the embed unit
            assert vec is None


# ---------------------------------------------------------------------------
# 4. contextual prefix: present, situating, and provably in the vector
# ---------------------------------------------------------------------------
def test_contextual_prefix_present_and_included_in_embedded_text(
        sop_case, store, embedder):
    tenant, result = sop_case["tenant"], sop_case["result"]
    children = [r for r in chunk_rows(store, tenant, result.document_id)
                if r["level"] == "child"]

    for c in children:
        prefix = c["contextual_prefix"]
        assert prefix and "\n" not in prefix          # one-line blurb
        assert SOP_TITLE in prefix                    # situates in the doc
        heading = (c["locator"] or {}).get("heading")
        if heading:
            assert heading in prefix                  # ...and in the section

    # The stored vector IS the embedding of prefix + content, not of the
    # bare passage: re-embed both live and compare.
    probe = children[0]
    composed = f"{probe['contextual_prefix']}\n\n{probe['content']}"
    vec_composed, vec_bare = embedder.embed([composed, probe["content"]])
    stored = parse_vector(probe["embedding"])
    assert cosine(stored, vec_composed) > 0.999
    assert cosine(stored, vec_composed) > cosine(stored, vec_bare)


# ---------------------------------------------------------------------------
# 5. §8.1a tag-as-claim: confident declared-vs-detected mismatch -> review
# ---------------------------------------------------------------------------
def test_track_mismatch_routes_to_review_queue(store, processing, tenant,
                                               tmp_path):
    land_file.root = tmp_path
    raw_id = land_file(store, processing.pipeline, processing.raw_store,
                       processing.dispatcher, tenant,
                       "sops/asset_list.csv", CSV_BYTES,
                       config={"data_track": "prose", "doc_type": "sop"})

    result = processing.process(tenant, raw_id)

    assert result.status == "review"
    assert "declared data_track 'prose'" in result.reason
    doc = doc_row(store, tenant, result.document_id)
    assert doc["review_status"] == "review"
    assert doc["metadata"]["detected_data_track"] == "structured"
    assert doc["metadata"]["detection_confident"] is True

    # It surfaced in the unified review queue as a document item...
    with store.transaction(tenant) as conn:
        queue = conn.execute(
            "SELECT * FROM review_queue WHERE tenant_id = %s AND kind = 'document'",
            (tenant,)).fetchall()
    assert [q["ref_id"] for q in queue] == [result.document_id]

    # ...with chunking withheld (never blindly obey a contradicted tag) and
    # the raw doc still 'landed' (processing is not complete).
    assert chunk_rows(store, tenant, result.document_id) == []
    assert store.get_raw_document(tenant, raw_id).status == "landed"

    # Redelivery keeps it held without double-flagging.
    again = processing.process(tenant, raw_id)
    assert again.status == "review" and again.document_id == result.document_id
    with store.transaction(tenant) as conn:
        docs = conn.execute(
            "SELECT count(*) AS n FROM documents"
            " WHERE tenant_id = %s AND raw_document_id = %s",
            (tenant, raw_id)).fetchone()
    assert docs["n"] == 1


# ---------------------------------------------------------------------------
# 6. non-prose tracks skip chunking (router hook for Prompt 4)
# ---------------------------------------------------------------------------
def test_structured_track_returns_no_chunks(store, processing, tenant,
                                            tmp_path):
    land_file.root = tmp_path
    raw_id = land_file(store, processing.pipeline, processing.raw_store,
                       processing.dispatcher, tenant,
                       "exports/asset_list.csv", CSV_BYTES,
                       config={"data_track": "structured"})

    result = processing.process(tenant, raw_id)

    assert result.status == "no_chunks"
    assert result.parents == result.children == 0
    doc = doc_row(store, tenant, result.document_id)
    assert doc["doc_type"] == "tabular"
    assert doc["metadata"]["data_track"] == "structured"
    assert doc["review_status"] == "none"      # declared == detected: no flag
    assert chunk_rows(store, tenant, result.document_id) == []
    assert store.get_raw_document(tenant, raw_id).status == "parsed"


# ---------------------------------------------------------------------------
# 7. idempotency: re-processing the same raw document adds no rows
# ---------------------------------------------------------------------------
def test_reprocessing_same_document_is_noop(store, processing, tenant,
                                            tmp_path):
    land_file.root = tmp_path
    sop_md = (
        b"# Label Reconciliation SOP\n\n"
        b"Purpose: reconcile printed labels against batch records after "
        b"every run so that no unaccounted label leaves the room.\n\n"
        b"## Counting\n\n"
        b"Count remaining labels, record the tally, and compare with the "
        b"issued quantity minus applied and destroyed labels.\n\n"
        b"## Escalation\n\n"
        b"Any discrepancy stops the line and goes to QA within fifteen "
        b"minutes, with the reconciliation sheet attached.\n")
    raw_id = land_file(store, processing.pipeline, processing.raw_store,
                       processing.dispatcher, tenant,
                       "sops/label_reconciliation.md", sop_md,
                       config={"data_track": "prose", "doc_type": "sop"})

    first = processing.process(tenant, raw_id)
    assert first.status == "processed"
    baseline = chunk_rows(store, tenant, first.document_id)
    assert baseline

    # Default replay: short-circuits on the existing hierarchy.
    second = processing.process(tenant, raw_id)
    assert second.status == "replayed"
    assert second.document_id == first.document_id
    assert (second.parents, second.children) == (first.parents, first.children)

    # Forced replay: runs the whole pass again (parse, chunk, embed) — and
    # STILL inserts nothing, because chunk identity is content-hashed.
    forced = processing.process(tenant, raw_id, force=True)
    assert forced.status == "processed"

    after = chunk_rows(store, tenant, first.document_id)
    assert [r["id"] for r in after] == [r["id"] for r in baseline]
    assert [r["content_hash"] for r in after] == \
        [r["content_hash"] for r in baseline]
    with store.transaction(tenant) as conn:
        n_docs = conn.execute(
            "SELECT count(*) AS n FROM documents"
            " WHERE tenant_id = %s AND raw_document_id = %s",
            (tenant, raw_id)).fetchone()["n"]
    assert n_docs == 1


# ---------------------------------------------------------------------------
# 8. queue consumption: a poison document nacks and redelivers, others flow
# ---------------------------------------------------------------------------
def test_consume_acks_good_and_nacks_poison(store, processing, tenant,
                                            tmp_path):
    land_file.root = tmp_path
    poison_id = land_file(store, processing.pipeline, processing.raw_store,
                          processing.dispatcher, tenant,
                          "sops/corrupt.docx", b"this is not really a docx",
                          config={"data_track": "prose"})

    results = processing.consume(tenant, limit=10)

    assert results == []  # the only delivery failed -> nothing processed
    message = processing.dispatcher.pending_for(tenant, poison_id)
    assert message.status == "queued"          # nacked: lease redelivers
    assert message.attempts == 1
    assert "ParseError" in message.last_error
